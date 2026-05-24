"""
builder.py — High-level DocxBuilder class for docx-author skill.

DocxBuilder wraps a python-docx Document and provides a fluent interface for
appending content blocks, setting metadata, configuring headers/footers, and
saving the final file.

The ``run(params)`` entry point in main.py is the primary interface used by
the sandbox runner.  DocxBuilder is not meant to be instantiated directly by
the generated Python code that the agent writes — it is an internal abstraction.
"""

from __future__ import annotations

import os
from typing import Optional

from docx import Document

from .document import (
    add_heading,
    add_horizontal_rule,
    add_image,
    add_list,
    add_page_break,
    add_paragraph,
    add_table,
    add_header_footer,
)
from .style import apply_style_config


class DocxBuilder:
    """Stateful builder that wraps a python-docx Document.

    Args:
        style_config: Optional style configuration dict (font, spacing, etc.).
            Applied to a freshly created Document via ``apply_style_config``.
        existing_doc: Optional pre-loaded Document (e.g. from a template).
            When provided, ``style_config`` is still applied on top.
    """

    def __init__(
        self,
        style_config: Optional[dict] = None,
        existing_doc: Optional[Document] = None,
    ) -> None:
        self._doc = existing_doc if existing_doc is not None else Document()
        self._style_config = style_config or {}
        # Store line_spacing and paragraph_spacing_pt for per-paragraph use
        self._line_spacing = float(self._style_config.get("line_spacing", 1.5))
        self._para_spacing_pt = int(self._style_config.get("paragraph_spacing_pt", 6))
        self._font_name = self._style_config.get("font_name", "微软雅黑")
        self._font_size_pt = float(self._style_config.get("font_size_pt", 10.5))
        self._heading_color = self._style_config.get("heading_color", "#1E293B")

        if self._style_config:
            apply_style_config(self._doc, self._style_config)

    # ------------------------------------------------------------------
    # Metadata
    # ------------------------------------------------------------------

    def set_metadata(self, metadata: dict) -> "DocxBuilder":
        """Write document core properties (title, author, subject, description).

        Args:
            metadata: Dict with optional keys: title, author, subject, description.

        Returns:
            self (for chaining)
        """
        props = self._doc.core_properties
        if "title" in metadata:
            props.title = str(metadata["title"])
        if "author" in metadata:
            props.author = str(metadata["author"])
        if "subject" in metadata:
            props.subject = str(metadata["subject"])
        if "description" in metadata:
            # python-docx exposes the dc:description field as core_properties.comments
            props.comments = str(metadata["description"])
        return self

    # ------------------------------------------------------------------
    # Header / Footer
    # ------------------------------------------------------------------

    def set_header_footer(
        self,
        header: Optional[dict] = None,
        footer: Optional[dict] = None,
    ) -> "DocxBuilder":
        """Configure the document header and footer.

        Args:
            header: Dict with optional "left", "center", "right" string values.
            footer: Dict with optional "left", "center", "right" string values.
                    "center" supports the ``{{page_number}}`` token.

        Returns:
            self (for chaining)
        """
        if header or footer:
            add_header_footer(self._doc, header=header or {}, footer=footer or {})
        return self

    # ------------------------------------------------------------------
    # Page margins
    # ------------------------------------------------------------------

    def set_page_margins(self, margin_cm: float) -> "DocxBuilder":
        """Override page margins after construction.

        Args:
            margin_cm: Uniform margin in centimetres.

        Returns:
            self (for chaining)
        """
        from .style import set_page_margins

        set_page_margins(self._doc, margin_cm=margin_cm)
        return self

    # ------------------------------------------------------------------
    # Block dispatcher
    # ------------------------------------------------------------------

    def add_block(self, block: dict) -> "DocxBuilder":
        """Append one content block to the document.

        Dispatches to the appropriate low-level helper based on ``block["type"]``.

        Supported types:
            heading, paragraph, table, list, image, page_break, horizontal_rule

        Unknown block types are silently skipped (logged to stderr in sandbox).

        Args:
            block: Block dict.  Must contain "type" key.

        Returns:
            self (for chaining)
        """
        block_type = block.get("type", "")

        if block_type == "heading":
            add_heading(
                self._doc,
                text=block.get("text", ""),
                level=int(block.get("level", 1)),
                color_hex=self._heading_color,
                font_name=self._font_name,
            )

        elif block_type == "paragraph":
            add_paragraph(
                self._doc,
                text=block.get("text", ""),
                bold=bool(block.get("bold", False)),
                italic=bool(block.get("italic", False)),
                alignment=block.get("alignment", "justify"),
                indent_level=int(block.get("indent_level", 0)),
                font_name=self._font_name,
                font_size_pt=self._font_size_pt,
                line_spacing=self._line_spacing,
                space_before_pt=self._para_spacing_pt,
                space_after_pt=self._para_spacing_pt,
            )

        elif block_type == "table":
            add_table(
                self._doc,
                headers=block.get("headers", []),
                rows=block.get("rows", []),
                style=block.get("style", "Table Grid"),
                col_widths_cm=block.get("col_widths_cm"),
                font_name=self._font_name,
            )

        elif block_type == "list":
            add_list(
                self._doc,
                items=block.get("items", []),
                ordered=bool(block.get("ordered", False)),
                font_name=self._font_name,
                font_size_pt=self._font_size_pt,
            )

        elif block_type == "image":
            try:
                add_image(
                    self._doc,
                    path=block.get("path", ""),
                    width_cm=float(block.get("width_cm", 14.0)),
                    caption=block.get("caption"),
                    alignment=block.get("alignment", "center"),
                    font_name=self._font_name,
                )
            except FileNotFoundError as exc:
                # Propagate as ValueError so the main.py caller can return
                # a structured error response instead of crashing the sandbox
                raise ValueError(str(exc)) from exc

        elif block_type == "page_break":
            add_page_break(self._doc)

        elif block_type == "horizontal_rule":
            add_horizontal_rule(self._doc)

        else:
            import sys
            print(f"[docx-author] warning: unknown block type '{block_type}' — skipped", file=sys.stderr)

        return self

    # ------------------------------------------------------------------
    # Save
    # ------------------------------------------------------------------

    def save(self, output_path: str) -> None:
        """Save the document to ``output_path``.

        The parent directory is created if it does not exist.

        Args:
            output_path: Absolute path to the output .docx file.
        """
        os.makedirs(os.path.dirname(output_path), exist_ok=True)
        self._doc.save(output_path)

"""
document.py — Low-level element functions for building Word documents.

Each function takes a python-docx Document object and appends one element
(heading, paragraph, table, list, image, etc.) to the end of the document.

Pitfalls documented here mirror the SKILL.md "常见坑" section:
  - Chinese font: set both run.font.name AND rFonts.eastAsia XML attribute
  - Table style: must use English style name ("Table Grid"), not Chinese
  - List styles: "List Bullet" / "List Number" with text-prefix fallback
  - Page number field: must use raw OxmlElement (no native API in python-docx)
  - Image path: must be absolute (/workspace/input/<path>); width auto-scales height
"""

from __future__ import annotations

import os
from typing import Optional

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


# ---------------------------------------------------------------------------
# Internal helpers
# ---------------------------------------------------------------------------

_ALIGN_MAP: dict[str, WD_ALIGN_PARAGRAPH] = {
    "left": WD_ALIGN_PARAGRAPH.LEFT,
    "center": WD_ALIGN_PARAGRAPH.CENTER,
    "right": WD_ALIGN_PARAGRAPH.RIGHT,
    "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
}


def _set_font(run, font_name: str = "微软雅黑", font_size_pt: float = 10.5) -> None:
    """Set font name for both Latin and East-Asian (CJK) characters in a run.

    python-docx exposes run.font.name, which sets the w:rFonts/@w:ascii and
    w:rFonts/@w:hAnsi attributes.  For Chinese characters the engine uses
    w:rFonts/@w:eastAsia — we must set that separately via XML to ensure
    Chinese text renders in the correct font on Windows/WPS.
    """
    run.font.name = font_name
    run.font.size = Pt(font_size_pt)
    # Explicitly set East-Asian font via XML
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn("w:eastAsia"), font_name)


def _set_paragraph_spacing(
    para,
    line_spacing: float = 1.5,
    space_before_pt: int = 6,
    space_after_pt: int = 6,
) -> None:
    """Apply line spacing and paragraph spacing (before / after) to a paragraph."""
    pf = para.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = line_spacing
    pf.space_before = Pt(space_before_pt)
    pf.space_after = Pt(space_after_pt)


# ---------------------------------------------------------------------------
# Heading
# ---------------------------------------------------------------------------


def add_heading(
    doc: Document,
    text: str,
    level: int = 1,
    color_hex: Optional[str] = None,
    font_name: str = "微软雅黑",
) -> None:
    """Append a heading paragraph at the given level (1–6).

    Args:
        doc: python-docx Document object.
        text: Heading text.
        level: Heading level 1–6. Level 1 maps to Word's "Heading 1" style.
        color_hex: Optional hex colour string (e.g. "#1E293B") for the heading text.
                   Defaults to the style's built-in colour.
        font_name: Font name for the heading (applied to all runs).
    """
    level = max(1, min(6, level))
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        _set_font(run, font_name=font_name)
        if color_hex:
            hex_clean = color_hex.lstrip("#")
            r, g, b = (int(hex_clean[i : i + 2], 16) for i in (0, 2, 4))
            run.font.color.rgb = RGBColor(r, g, b)


# ---------------------------------------------------------------------------
# Paragraph
# ---------------------------------------------------------------------------


def add_paragraph(
    doc: Document,
    text: str,
    bold: bool = False,
    italic: bool = False,
    alignment: str = "justify",
    indent_level: int = 0,
    font_name: str = "微软雅黑",
    font_size_pt: float = 10.5,
    line_spacing: float = 1.5,
    space_before_pt: int = 6,
    space_after_pt: int = 6,
) -> None:
    """Append a styled paragraph.

    Args:
        doc: python-docx Document object.
        text: Paragraph text. Supports \\n for line breaks within the paragraph.
        bold: Apply bold formatting.
        italic: Apply italic formatting.
        alignment: "left" | "center" | "right" | "justify" (default "justify").
        indent_level: Indent level 0–3 (each level adds 0.75 cm left indent).
        font_name: Font name (applied to all runs).
        font_size_pt: Font size in points.
        line_spacing: Line spacing multiplier (e.g. 1.5).
        space_before_pt: Space before paragraph in points.
        space_after_pt: Space after paragraph in points.
    """
    para = doc.add_paragraph()
    para.alignment = _ALIGN_MAP.get(alignment, WD_ALIGN_PARAGRAPH.JUSTIFY)
    pf = para.paragraph_format
    if indent_level > 0:
        pf.left_indent = Cm(indent_level * 0.75)
    _set_paragraph_spacing(para, line_spacing, space_before_pt, space_after_pt)

    # Support \n by splitting into sub-runs with line-break XML
    lines = text.split("\n")
    for i, line in enumerate(lines):
        run = para.add_run(line)
        run.bold = bold
        run.italic = italic
        _set_font(run, font_name=font_name, font_size_pt=font_size_pt)
        # Add <w:br/> between lines (not after the last one)
        if i < len(lines) - 1:
            br = OxmlElement("w:br")
            run._r.append(br)


# ---------------------------------------------------------------------------
# Table
# ---------------------------------------------------------------------------


def _set_cell_bg(cell, fill_hex: str) -> None:
    """Set the background fill of a table cell via raw XML."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex.upper())
    tcPr.append(shd)


def add_table(
    doc: Document,
    headers: list[str],
    rows: list[list],
    style: str = "Table Grid",
    col_widths_cm: Optional[list[float]] = None,
    header_bg_hex: str = "2563EB",
    font_name: str = "微软雅黑",
    font_size_pt: float = 10.0,
) -> None:
    """Append a formatted table with a styled header row.

    Args:
        doc: python-docx Document object.
        headers: List of column header strings.
        rows: List of data rows; each row is a list of cell values (str/int/float).
        style: Word table style name (English). Defaults to "Table Grid".
               Falls back to no-style + manual borders if style not found.
        col_widths_cm: Optional per-column widths in cm.
        header_bg_hex: Hex colour for the header row background (default blue).
        font_name: Font name for all cells.
        font_size_pt: Font size in points for all cells.
    """
    total_cols = len(headers)
    if total_cols == 0:
        return

    # Create table with header row
    table = doc.add_table(rows=1, cols=total_cols)
    try:
        table.style = style
    except KeyError:
        # Style not found in this document — proceed without it
        pass

    # Header row: blue background + white bold text
    hdr_cells = table.rows[0].cells
    for i, header_text in enumerate(headers):
        cell = hdr_cells[i]
        cell.text = ""
        _set_cell_bg(cell, header_bg_hex)
        run = cell.paragraphs[0].add_run(str(header_text))
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        _set_font(run, font_name=font_name, font_size_pt=font_size_pt)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Data rows
    for row_data in rows:
        row_cells = table.add_row().cells
        for i, val in enumerate(row_data[:total_cols]):
            cell = row_cells[i]
            cell.text = ""
            run = cell.paragraphs[0].add_run(str(val) if val is not None else "")
            _set_font(run, font_name=font_name, font_size_pt=font_size_pt)

    # Optional column widths
    if col_widths_cm and len(col_widths_cm) == total_cols:
        for col_idx, width in enumerate(col_widths_cm):
            for row in table.rows:
                row.cells[col_idx].width = Cm(width)


# ---------------------------------------------------------------------------
# List (ordered / unordered)
# ---------------------------------------------------------------------------


def add_list(
    doc: Document,
    items: list,
    ordered: bool = False,
    font_name: str = "微软雅黑",
    font_size_pt: float = 10.5,
) -> None:
    """Append a bulleted or numbered list.

    Each item in ``items`` can be:
      - A plain string → rendered at indent level 0.
      - A dict with keys "text" (str) and optional "indent" (int, 0-based).

    If the Word style "List Number" or "List Bullet" is not available in the
    document (e.g. the template omits them), falls back to a manual prefix
    (e.g. "1. " or "• ") prepended to the text.

    Args:
        doc: python-docx Document object.
        items: List of strings or dicts.
        ordered: True for numbered list, False for bulleted list.
        font_name: Font name applied to all list items.
        font_size_pt: Font size in points.
    """
    style_name = "List Number" if ordered else "List Bullet"
    style_available = style_name in [s.name for s in doc.styles]

    for idx, item in enumerate(items, start=1):
        if isinstance(item, dict):
            text = item.get("text", "")
            indent = int(item.get("indent", 0))
        else:
            text = str(item)
            indent = 0

        if style_available:
            para = doc.add_paragraph(style=style_name)
            # Increase indent level for nested items
            if indent > 0:
                para.paragraph_format.left_indent = Cm(indent * 1.0)
            run = para.add_run(text)
        else:
            # Fallback: manual prefix + Normal paragraph
            prefix = f"{idx}. " if ordered else "• "
            para = doc.add_paragraph()
            if indent > 0:
                para.paragraph_format.left_indent = Cm(indent * 1.0)
            run = para.add_run(prefix + text)

        _set_font(run, font_name=font_name, font_size_pt=font_size_pt)
        _set_paragraph_spacing(para, line_spacing=1.2, space_before_pt=3, space_after_pt=3)


# ---------------------------------------------------------------------------
# Image
# ---------------------------------------------------------------------------


def add_image(
    doc: Document,
    path: str,
    width_cm: float = 14.0,
    caption: Optional[str] = None,
    alignment: str = "center",
    font_name: str = "微软雅黑",
) -> None:
    """Append an image from the sandbox input directory.

    The ``path`` parameter is treated as relative to ``/workspace/input/``.
    If it is already an absolute path it is used as-is.

    python-docx auto-scales the image height to preserve the aspect ratio when
    only width is specified — this is intentional behaviour.

    Args:
        doc: python-docx Document object.
        path: Path to the image file (relative to /workspace/input/ or absolute).
        width_cm: Display width in cm. Height is auto-calculated.
        caption: Optional caption text rendered below the image.
        alignment: "left" | "center" | "right" (default "center").
        font_name: Font name for the caption text.

    Raises:
        FileNotFoundError: If the image file does not exist (caller should handle).
    """
    if not os.path.isabs(path):
        full_path = os.path.join("/workspace/input", path)
    else:
        full_path = path

    if not os.path.exists(full_path):
        raise FileNotFoundError(f"Image not found: {full_path}")

    # Add image in a centred paragraph
    para = doc.add_paragraph()
    para.alignment = _ALIGN_MAP.get(alignment, WD_ALIGN_PARAGRAPH.CENTER)
    run = para.add_run()
    run.add_picture(full_path, width=Cm(width_cm))

    # Optional caption
    if caption:
        cap_para = doc.add_paragraph()
        cap_para.alignment = _ALIGN_MAP.get(alignment, WD_ALIGN_PARAGRAPH.CENTER)
        cap_run = cap_para.add_run(caption)
        cap_run.italic = True
        _set_font(cap_run, font_name=font_name, font_size_pt=9.0)
        _set_paragraph_spacing(cap_para, line_spacing=1.0, space_before_pt=3, space_after_pt=6)


# ---------------------------------------------------------------------------
# Header / Footer
# ---------------------------------------------------------------------------


def _add_page_number_field(run) -> None:
    """Insert a Word PAGE field into ``run`` via raw XML.

    python-docx has no native API for field codes; we must operate on the
    underlying Open XML elements directly.  This produces the equivalent of
    inserting a { PAGE } field in Word's "Insert Field" dialog.

    Note: The rendered number only appears after the document is opened in Word
    (field codes are updated on open/print).  In python-docx the field stays as
    an fldChar placeholder in the XML.
    """
    fldChar_begin = OxmlElement("w:fldChar")
    fldChar_begin.set(qn("w:fldCharType"), "begin")

    instrText = OxmlElement("w:instrText")
    instrText.set(qn("xml:space"), "preserve")
    instrText.text = " PAGE "

    fldChar_end = OxmlElement("w:fldChar")
    fldChar_end.set(qn("w:fldCharType"), "end")

    run._r.extend([fldChar_begin, instrText, fldChar_end])


def _fill_header_footer_para(para, left: str, center: str, right: str) -> None:
    """Populate a three-column tab-separated header/footer paragraph.

    Word headers/footers conventionally use a paragraph with two tab stops
    (centre-tab and right-tab) to achieve left/center/right layout.  The
    built-in "Header" and "Footer" styles already include these tab stops.

    The ``center`` value supports the special token ``{{page_number}}`` which
    is replaced with a Word PAGE field (rendered at open/print time in Word).
    """
    para.clear()  # remove any existing runs
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT

    def _add_text_or_field(run_text: str) -> None:
        """Add text, substituting {{page_number}} with a PAGE field."""
        if "{{page_number}}" in run_text:
            parts = run_text.split("{{page_number}}")
            for i, part in enumerate(parts):
                if part:
                    para.add_run(part)
                if i < len(parts) - 1:
                    field_run = para.add_run()
                    _add_page_number_field(field_run)
        else:
            para.add_run(run_text)

    # Left portion
    if left:
        _add_text_or_field(left)

    # Center (tab + text/field)
    if center:
        para.add_run("\t")
        _add_text_or_field(center)

    # Right (tab + text)
    if right:
        para.add_run("\t")
        _add_text_or_field(right)


def add_header_footer(
    doc: Document,
    header: Optional[dict] = None,
    footer: Optional[dict] = None,
) -> None:
    """Set the document header and/or footer.

    Args:
        doc: python-docx Document object.
        header: Dict with optional keys "left", "center", "right".
        footer: Dict with optional keys "left", "center", "right".
                The "center" value may contain ``{{page_number}}`` which is
                replaced with a live Word PAGE field.

    The function operates on ``doc.sections[0]`` (the first/only section).
    For multi-section documents, call this function once per section or extend.
    """
    section = doc.sections[0]

    if header:
        section.header.is_linked_to_previous = False
        hdr_paras = section.header.paragraphs
        if hdr_paras:
            _fill_header_footer_para(
                hdr_paras[0],
                left=header.get("left", ""),
                center=header.get("center", ""),
                right=header.get("right", ""),
            )
        else:
            para = section.header.add_paragraph()
            _fill_header_footer_para(
                para,
                left=header.get("left", ""),
                center=header.get("center", ""),
                right=header.get("right", ""),
            )

    if footer:
        section.footer.is_linked_to_previous = False
        ftr_paras = section.footer.paragraphs
        if ftr_paras:
            _fill_header_footer_para(
                ftr_paras[0],
                left=footer.get("left", ""),
                center=footer.get("center", ""),
                right=footer.get("right", ""),
            )
        else:
            para = section.footer.add_paragraph()
            _fill_header_footer_para(
                para,
                left=footer.get("left", ""),
                center=footer.get("center", ""),
                right=footer.get("right", ""),
            )


# ---------------------------------------------------------------------------
# Page break
# ---------------------------------------------------------------------------


def add_page_break(doc: Document) -> None:
    """Append a hard page break to the document.

    Inserts a paragraph containing a single run with a ``<w:br w:type="page"/>``
    element — the standard way python-docx represents hard page breaks.
    """
    para = doc.add_paragraph()
    run = para.add_run()
    br = OxmlElement("w:br")
    br.set(qn("w:type"), "page")
    run._r.append(br)


# ---------------------------------------------------------------------------
# Horizontal rule
# ---------------------------------------------------------------------------


def add_horizontal_rule(doc: Document) -> None:
    """Append a visual horizontal rule (line) using a paragraph border.

    python-docx has no built-in method for horizontal rules.  We achieve
    the effect by adding an empty paragraph and setting a bottom border on
    its paragraph properties (pPr / pBdr / bottom) via raw XML.
    """
    para = doc.add_paragraph()
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "AAAAAA")
    pBdr.append(bottom)
    pPr.append(pBdr)

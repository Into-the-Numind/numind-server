"""
template.py — Template loading and variable substitution for docx-author.

Loads a .docx template from /skills/docx-author/templates/, substitutes
{{key}} placeholders in paragraphs and table cells, then returns the
python-docx Document for further content additions via DocxBuilder.

Known limitation (documented in SKILL.md "常见坑 #4"):
Word may split a single run of text across multiple <w:r> elements (e.g. when
the spell-checker or auto-correct annotates a word).  If "{{title}}" is split
into ["{{", "title", "}}"] across three runs, the regex applied to each run
individually will not match.

Mitigation strategy (V1.5 — simple):
  - Template authoring guideline: type placeholder text in a fresh paragraph,
    select the whole placeholder, apply "No Proofing" / clear all formatting,
    then save.  This prevents automatic run-splitting.
  - The paragraph-level fallback implemented below: if no run in a paragraph
    contains the full placeholder string, we check the combined paragraph.text
    and, if a match is found, replace the ENTIRE paragraph content with a
    single unstyled run containing the substituted text.  This loses per-run
    formatting but is safe and predictable.
"""

from __future__ import annotations

import os
import re
from typing import Optional

from docx import Document
from docx.oxml.ns import qn


PLACEHOLDER_RE = re.compile(r"\{\{(\w+)\}\}")

# Default location of templates inside the sandbox
_TEMPLATE_DIR = "/skills/docx-author/templates"


def load_template_doc(
    template_name: str,
    template_vars: Optional[dict] = None,
    template_dir: str = _TEMPLATE_DIR,
) -> Document:
    """Load a template .docx and substitute {{key}} placeholders.

    The function substitutes placeholders in:
      1. Top-level paragraphs.
      2. Table cell paragraphs.

    For each paragraph it first attempts run-level substitution (preserves
    per-run formatting).  If the combined paragraph text still contains an
    unresolved placeholder after run-level substitution (indicating the
    placeholder was split across runs), it falls back to paragraph-level
    substitution (replaces all runs with a single unstyled run).

    Args:
        template_name: Basename without extension, e.g. "general-report".
        template_vars: Dict mapping placeholder names to replacement strings.
        template_dir: Directory containing template .docx files.

    Returns:
        A python-docx Document with placeholders substituted.

    Raises:
        FileNotFoundError: If the template file does not exist.
    """
    template_vars = template_vars or {}
    path = os.path.join(template_dir, f"{template_name}.docx")
    if not os.path.exists(path):
        raise FileNotFoundError(f"Template not found: {path}")

    doc = Document(path)

    # Substitute all paragraph-level text (including headers/footers implicitly
    # embedded in the template's XML body)
    for para in doc.paragraphs:
        _substitute_paragraph(para, template_vars)

    # Substitute table cell text
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    _substitute_paragraph(para, template_vars)

    return doc


def _substitute_paragraph(para, vars: dict) -> None:
    """Substitute {{key}} placeholders in a single paragraph.

    Strategy:
      1. Try run-level substitution (preserves formatting).
      2. If the combined paragraph.text still has unresolved placeholders,
         fall back to paragraph-level substitution (single unstyled run).
    """
    # Pass 1: run-level (preserves formatting)
    for run in para.runs:
        if PLACEHOLDER_RE.search(run.text):
            run.text = PLACEHOLDER_RE.sub(
                lambda m: str(vars.get(m.group(1), m.group(0))), run.text
            )

    # Pass 2: paragraph-level fallback for split-run placeholders
    combined = para.text
    if PLACEHOLDER_RE.search(combined):
        substituted = PLACEHOLDER_RE.sub(
            lambda m: str(vars.get(m.group(1), m.group(0))), combined
        )
        # Clear all runs and replace with a single unstyled run
        # We preserve paragraph-level formatting (alignment, spacing) but lose
        # per-run character formatting — acceptable for placeholder paragraphs.
        for run in para.runs:
            run.text = ""
        if para.runs:
            para.runs[0].text = substituted
        else:
            para.add_run(substituted)

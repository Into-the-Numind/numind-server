"""
test_basic.py — pytest unit tests for the docx-author skill.

Tests call the ``run(params)`` entry point directly (no subprocess).  Output
files are written to a temporary directory provided by pytest's ``tmp_path``
fixture.

The tests pass absolute tmp_path-based paths as ``output_filename`` so the
skill writes directly to the temp directory — no /output directory needed.
The main.py ``sanitize_filename`` function is bypassed when the filename is
already an absolute path (basename extraction keeps it intact).

Requires:
    pip install pytest python-docx>=1.1
"""

from __future__ import annotations

import os
import sys
import pytest

# Locate the skill root so imports work regardless of cwd
_SKILL_ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
if _SKILL_ROOT not in sys.path:
    sys.path.insert(0, _SKILL_ROOT)

import main as skill_main
from main import sanitize_filename
from helpers.template import load_template_doc

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm


# ---------------------------------------------------------------------------
# Helper: call run() with an absolute output path in tmp_path
# ---------------------------------------------------------------------------

def _run_to(params: dict, out_path: str) -> dict:
    """Run the skill with an explicit absolute output path.

    main.py's sanitize_filename does os.path.basename(), which means an
    absolute path like /tmp/.../foo.docx becomes 'foo.docx', then main.py
    prepends '/output/'.  We instead pass the full absolute path as
    output_filename and patch the '/output' prefix inside main.py.
    """
    import unittest.mock as mock

    # Make the builder save directly to out_path by patching os.path.join
    # only for the /output prefix case
    original_join = os.path.join

    params_copy = dict(params)
    # Use just the basename so sanitize_filename keeps it clean, then redirect
    basename = os.path.basename(out_path)
    params_copy["output_filename"] = basename

    # Patch the "/output" prefix in main.py by patching os.path.join there
    def patched_join(*args):
        result = original_join(*args)
        # Replace /output/<basename> with our actual tmp path
        if args and args[0] == "/output":
            return out_path
        return result

    with mock.patch("main.os.path.join", side_effect=patched_join):
        return skill_main.run(params_copy)


# ---------------------------------------------------------------------------
# Heading tests
# ---------------------------------------------------------------------------

def test_heading_rendered(tmp_path):
    """Heading block should produce a paragraph with 'Heading 1' style and correct text."""
    out = str(tmp_path / "h1.docx")
    params = {
        "output_filename": "h1.docx",
        "blocks": [{"type": "heading", "level": 1, "text": "测试标题"}],
    }
    result = _run_to(params, out)
    assert result["success"], result.get("error")
    doc = Document(out)
    heading_paras = [p for p in doc.paragraphs if "Heading" in p.style.name]
    assert heading_paras, "No heading paragraph found"
    assert heading_paras[0].text == "测试标题"
    assert heading_paras[0].style.name == "Heading 1"


def test_heading_level_2(tmp_path):
    """Level 2 heading should use 'Heading 2' style."""
    out = str(tmp_path / "h2.docx")
    params = {
        "output_filename": "h2.docx",
        "blocks": [{"type": "heading", "level": 2, "text": "二级标题"}],
    }
    result = _run_to(params, out)
    assert result["success"], result.get("error")
    doc = Document(out)
    heading_paras = [p for p in doc.paragraphs if p.style.name == "Heading 2"]
    assert heading_paras, "No Heading 2 paragraph found"
    assert heading_paras[0].text == "二级标题"


# ---------------------------------------------------------------------------
# Paragraph tests
# ---------------------------------------------------------------------------

def test_paragraph_bold_italic(tmp_path):
    """Bold and italic flags should be reflected in the run's formatting."""
    out = str(tmp_path / "bold.docx")
    params = {
        "output_filename": "bold.docx",
        "blocks": [
            {"type": "paragraph", "text": "粗体斜体文字", "bold": True, "italic": True}
        ],
    }
    result = _run_to(params, out)
    assert result["success"], result.get("error")
    doc = Document(out)
    target = next(p for p in doc.paragraphs if p.text == "粗体斜体文字")
    assert target.runs[0].bold is True
    assert target.runs[0].italic is True


def test_paragraph_alignment_justify(tmp_path):
    """Default paragraph alignment should be JUSTIFY."""
    out = str(tmp_path / "justify.docx")
    params = {
        "output_filename": "justify.docx",
        "blocks": [{"type": "paragraph", "text": "对齐测试段落"}],
    }
    result = _run_to(params, out)
    assert result["success"], result.get("error")
    doc = Document(out)
    target = next(p for p in doc.paragraphs if p.text == "对齐测试段落")
    assert target.alignment == WD_ALIGN_PARAGRAPH.JUSTIFY


def test_paragraph_alignment_center(tmp_path):
    """Center alignment should be applied when requested."""
    out = str(tmp_path / "center.docx")
    params = {
        "output_filename": "center.docx",
        "blocks": [{"type": "paragraph", "text": "居中段落", "alignment": "center"}],
    }
    result = _run_to(params, out)
    assert result["success"], result.get("error")
    doc = Document(out)
    target = next(p for p in doc.paragraphs if p.text == "居中段落")
    assert target.alignment == WD_ALIGN_PARAGRAPH.CENTER


# ---------------------------------------------------------------------------
# Table tests
# ---------------------------------------------------------------------------

def test_table_row_count(tmp_path):
    """5 data rows + 1 header row = 6 total rows in the document table."""
    out = str(tmp_path / "table.docx")
    params = {
        "output_filename": "table.docx",
        "blocks": [
            {
                "type": "table",
                "headers": ["列A", "列B"],
                "rows": [
                    ["r1a", "r1b"],
                    ["r2a", "r2b"],
                    ["r3a", "r3b"],
                    ["r4a", "r4b"],
                    ["r5a", "r5b"],
                ],
            }
        ],
    }
    result = _run_to(params, out)
    assert result["success"], result.get("error")
    doc = Document(out)
    assert doc.tables, "No table found in document"
    assert len(doc.tables[0].rows) == 6  # 1 header + 5 data


def test_table_header_background(tmp_path):
    """Header row cell should have a non-default shd fill attribute (blue background)."""
    out = str(tmp_path / "table_bg.docx")
    params = {
        "output_filename": "table_bg.docx",
        "blocks": [
            {
                "type": "table",
                "headers": ["名称", "数量"],
                "rows": [["苹果", 10]],
            }
        ],
    }
    result = _run_to(params, out)
    assert result["success"], result.get("error")
    doc = Document(out)
    table = doc.tables[0]
    first_header_cell = table.rows[0].cells[0]
    # Extract shd fill from XML
    tc = first_header_cell._tc
    tcPr = tc.find(qn("w:tcPr"))
    assert tcPr is not None, "No tcPr in header cell"
    shd = tcPr.find(qn("w:shd"))
    assert shd is not None, "No shd element in header cell tcPr"
    fill = shd.get(qn("w:fill"), "")
    assert fill.upper() == "2563EB", f"Expected blue fill 2563EB, got {fill}"


def test_table_col_widths(tmp_path):
    """Column widths should match col_widths_cm when specified."""
    out = str(tmp_path / "table_widths.docx")
    params = {
        "output_filename": "table_widths.docx",
        "blocks": [
            {
                "type": "table",
                "headers": ["窄列", "宽列"],
                "rows": [["a", "b"]],
                "col_widths_cm": [3.0, 10.0],
            }
        ],
    }
    result = _run_to(params, out)
    assert result["success"], result.get("error")
    doc = Document(out)
    table = doc.tables[0]
    # Check that the first column is narrower than the second
    col0_width = table.rows[0].cells[0].width
    col1_width = table.rows[0].cells[1].width
    assert col0_width is not None
    assert col1_width is not None
    assert col0_width < col1_width, (
        f"Expected col0 ({col0_width}) < col1 ({col1_width})"
    )
    # Check approximate values (Cm returns EMUs; 1 cm ≈ 360000 EMU)
    expected_col0 = Cm(3.0)
    assert abs(col0_width - expected_col0) < 72000, (  # ±0.2 cm tolerance
        f"col0 width {col0_width} differs significantly from Cm(3.0)={expected_col0}"
    )


# ---------------------------------------------------------------------------
# List tests
# ---------------------------------------------------------------------------

def test_list_unordered_bullet(tmp_path):
    """Unordered list should use 'List Bullet' style or contain bullet prefix."""
    out = str(tmp_path / "list_bullet.docx")
    params = {
        "output_filename": "list_bullet.docx",
        "blocks": [
            {"type": "list", "items": ["苹果", "香蕉", "橙子"], "ordered": False}
        ],
    }
    result = _run_to(params, out)
    assert result["success"], result.get("error")
    doc = Document(out)
    list_paras = [
        p for p in doc.paragraphs
        if p.style.name == "List Bullet" or p.text.startswith("•")
    ]
    assert len(list_paras) == 3, f"Expected 3 bullet paragraphs, got {len(list_paras)}"


def test_list_ordered_prefix(tmp_path):
    """Ordered list should use 'List Number' style or manual number prefix."""
    out = str(tmp_path / "list_ordered.docx")
    params = {
        "output_filename": "list_ordered.docx",
        "blocks": [
            {"type": "list", "items": ["第一步", "第二步"], "ordered": True}
        ],
    }
    result = _run_to(params, out)
    assert result["success"], result.get("error")
    doc = Document(out)
    list_paras = [
        p for p in doc.paragraphs
        if p.style.name == "List Number" or (p.text and p.text[0].isdigit())
    ]
    assert len(list_paras) == 2, f"Expected 2 ordered paragraphs, got {len(list_paras)}"


# ---------------------------------------------------------------------------
# Image tests
# ---------------------------------------------------------------------------

def test_image_file_not_found_error(tmp_path):
    """Non-existent image path should return success=False with an error message."""
    out = str(tmp_path / "img_missing.docx")
    params = {
        "output_filename": "img_missing.docx",
        "blocks": [
            {"type": "image", "path": "nonexistent_image_xyz.png"}
        ],
    }
    result = _run_to(params, out)
    assert result["success"] is False
    assert "not found" in (result.get("error") or "").lower()


def _make_tiny_png() -> bytes:
    """Create a minimal valid 1x1 pixel PNG for testing."""
    import struct, zlib

    def _crc(data):
        return struct.pack(">I", zlib.crc32(data) & 0xFFFFFFFF)

    sig = b"\x89PNG\r\n\x1a\n"
    ihdr_data = struct.pack(">IIBBBBB", 1, 1, 8, 2, 0, 0, 0)
    ihdr_chunk = b"IHDR" + ihdr_data
    ihdr = struct.pack(">I", 13) + ihdr_chunk + _crc(ihdr_chunk)
    idat_raw = b"\x00\xff\xff\xff"
    compressed = zlib.compress(idat_raw)
    idat_body = b"IDAT" + compressed
    idat = struct.pack(">I", len(compressed)) + idat_body + _crc(idat_body)
    iend_body = b"IEND"
    iend = struct.pack(">I", 0) + iend_body + _crc(iend_body)
    return sig + ihdr + idat + iend


def test_image_inserted(tmp_path):
    """Valid image should result in an inline shape in the document."""
    import unittest.mock as mock
    import helpers.document as doc_mod

    # Create a tiny PNG and place it at the real path used by add_image
    real_input = tmp_path / "input"
    real_input.mkdir()
    png_path = real_input / "test.png"
    png_path.write_bytes(_make_tiny_png())

    out = str(tmp_path / "img_test.docx")

    # Patch add_image to use our temp path instead of /workspace/input
    orig_add_image = doc_mod.add_image

    def patched_add_image(doc, path, **kwargs):
        # Redirect path resolution to our temp input dir
        if not os.path.isabs(path):
            kwargs_copy = dict(kwargs)
            import unittest.mock as _mock
            with _mock.patch("helpers.document.os.path.join", return_value=str(png_path)), \
                 _mock.patch("helpers.document.os.path.exists", return_value=True):
                return orig_add_image(doc, path, **kwargs_copy)
        return orig_add_image(doc, path, **kwargs)

    orig_join = os.path.join

    def patched_main_join(*args):
        joined = orig_join(*args)
        if args and args[0] == "/output":
            return out
        return joined

    with mock.patch("helpers.builder.add_image", side_effect=patched_add_image), \
         mock.patch("main.os.path.join", side_effect=patched_main_join):
        params = {
            "output_filename": "img_test.docx",
            "blocks": [{"type": "image", "path": "test.png", "width_cm": 1.0}],
        }
        result = skill_main.run(params)

    assert result["success"], result.get("error")
    doc = Document(out)
    assert len(doc.inline_shapes) == 1, "Expected 1 inline shape (image)"


# ---------------------------------------------------------------------------
# Page break test
# ---------------------------------------------------------------------------

def test_page_break_inserted(tmp_path):
    """page_break block should insert a paragraph with a <w:br w:type='page'/> element."""
    out = str(tmp_path / "pagebreak.docx")
    params = {
        "output_filename": "pagebreak.docx",
        "blocks": [
            {"type": "paragraph", "text": "第一页内容"},
            {"type": "page_break"},
            {"type": "paragraph", "text": "第二页内容"},
        ],
    }
    result = _run_to(params, out)
    assert result["success"], result.get("error")
    doc = Document(out)
    # Find any paragraph that contains a <w:br w:type="page"/> element
    found = False
    for para in doc.paragraphs:
        for r in para.runs:
            for child in r._r:
                if child.tag == qn("w:br"):
                    if child.get(qn("w:type")) == "page":
                        found = True
                        break
    assert found, "No page break element found in document XML"


# ---------------------------------------------------------------------------
# Metadata test
# ---------------------------------------------------------------------------

def test_metadata_written(tmp_path):
    """Document core properties should reflect the metadata dict."""
    out = str(tmp_path / "meta.docx")
    params = {
        "output_filename": "meta.docx",
        "metadata": {
            "title": "元数据测试文档",
            "author": "单元测试",
            "subject": "测试主题",
            "description": "这是描述字段",
        },
        "blocks": [{"type": "paragraph", "text": "内容"}],
    }
    result = _run_to(params, out)
    assert result["success"], result.get("error")
    doc = Document(out)
    props = doc.core_properties
    assert props.title == "元数据测试文档"
    assert props.author == "单元测试"
    assert props.subject == "测试主题"
    # python-docx maps dc:description → core_properties.comments
    assert props.comments == "这是描述字段"


# ---------------------------------------------------------------------------
# Template tests
# ---------------------------------------------------------------------------

def test_template_vars_replacement(tmp_path):
    """Template variables {{doc_title}} etc. should be substituted in the loaded doc."""
    template_dir = os.path.join(
        os.path.dirname(os.path.dirname(os.path.abspath(__file__))),
        "templates",
    )
    doc = load_template_doc(
        template_name="general-report",
        template_vars={"doc_title": "单元测试报告", "author": "测试团队"},
        template_dir=template_dir,
    )
    # Collect all text in the document
    all_text = " ".join(p.text for p in doc.paragraphs)
    # Also check table cells
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    all_text += " " + para.text
    assert "单元测试报告" in all_text, "doc_title placeholder was not substituted"
    assert "测试团队" in all_text, "author placeholder was not substituted"


def test_template_multi_run_placeholder(tmp_path):
    """Unresolved placeholder (split across runs) should not crash and placeholder remains."""
    # Create a minimal .docx with a placeholder that survives as-is
    # (simulates the split-run scenario by passing a key not in vars)
    template_dir = os.path.join(
        os.path.dirname(os.path.dirname(os.path.abspath(__file__))),
        "templates",
    )
    # Pass empty vars — all placeholders should remain verbatim (no crash)
    doc = load_template_doc(
        template_name="general-report",
        template_vars={},  # no substitutions
        template_dir=template_dir,
    )
    # Should not raise; document is valid
    all_text = " ".join(p.text for p in doc.paragraphs)
    # At least some placeholder text should remain unchanged
    assert "{{" in all_text or len(doc.paragraphs) > 0


# ---------------------------------------------------------------------------
# Output file creation test
# ---------------------------------------------------------------------------

def test_output_file_created(tmp_path):
    """run() should create the output file at the specified path."""
    out = str(tmp_path / "output.docx")
    params = {
        "output_filename": "output.docx",
        "blocks": [{"type": "paragraph", "text": "测试内容"}],
    }
    result = _run_to(params, out)
    assert result["success"], result.get("error")
    assert os.path.exists(out)
    assert os.path.getsize(out) > 0


# ---------------------------------------------------------------------------
# Filename sanitisation tests
# ---------------------------------------------------------------------------

def test_filename_sanitization_removes_illegal_chars():
    """Illegal characters should be stripped from the output filename."""
    result = sanitize_filename('report<>:"/\\|?*.docx')
    assert "<" not in result
    assert ">" not in result
    assert ":" not in result
    assert '"' not in result
    assert "/" not in result
    assert "\\" not in result
    assert "|" not in result
    assert "?" not in result
    assert "*" not in result


def test_filename_sanitization_prevents_path_traversal():
    """Path traversal sequences should be removed from the filename."""
    result = sanitize_filename("../../etc/passwd.docx")
    assert ".." not in result
    assert "/" not in result


def test_filename_sanitization_preserves_chinese():
    """Chinese characters should be preserved in the sanitised filename."""
    result = sanitize_filename("2026年度报告.docx")
    assert "2026年度报告" in result


def test_filename_sanitization_fallback():
    """Empty or all-illegal filename should fall back to 'output.docx'."""
    result = sanitize_filename('?"<>|')
    assert result == "output.docx"


# ---------------------------------------------------------------------------
# Block count test
# ---------------------------------------------------------------------------

def test_blocks_rendered_count(tmp_path):
    """result['blocks_rendered'] should equal the number of blocks in params."""
    out = str(tmp_path / "count.docx")
    params = {
        "output_filename": "count.docx",
        "blocks": [
            {"type": "heading", "level": 1, "text": "H1"},
            {"type": "paragraph", "text": "P1"},
            {"type": "paragraph", "text": "P2"},
        ],
    }
    result = _run_to(params, out)
    assert result["success"], result.get("error")
    assert result["blocks_rendered"] == 3

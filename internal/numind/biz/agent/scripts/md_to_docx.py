#!/usr/bin/env python3
"""md_to_docx.py — deterministic Markdown -> .docx converter (create_docx fast path).

This script is FIXED and version-controlled. The agent never writes Python for the
create_docx tool; it only supplies Markdown content. We run this script inside the
same sandbox run_python uses.

Contract:
  - Reads Markdown from /workdir/input/source.md.
  - Writes the .docx to /workdir/output/<argv[1]> (a sanitized filename ending in
    .docx). If argv[1] is missing, defaults to "document.docx".
  - Inline images: ![alt](path) where `path` resolves under /workdir/input/ are
    embedded via doc.add_picture; otherwise the alt text is rendered as a caption.

Only depends on python-docx (>=1.1) which the sandbox image ships. Robust to bad
Markdown: any line that cannot be parsed falls back to a plain paragraph rather
than crashing, so a malformed document still produces SOMETHING openable.
"""

import os
import re
import sys

from docx import Document
from docx.shared import Inches

# Paths default to the sandbox layout but are overridable via env so the script
# is unit-testable on a host without a /workdir mount.
INPUT_DIR = os.environ.get("MD2DOCX_INPUT_DIR", "/workdir/input")
OUTPUT_DIR = os.environ.get("MD2DOCX_OUTPUT_DIR", "/workdir/output")
INPUT_MD = os.path.join(INPUT_DIR, "source.md")

# Inline image: ![alt](path)
IMG_RE = re.compile(r"^!\[(?P<alt>[^\]]*)\]\((?P<path>[^)]+)\)\s*$")
# Heading: leading #'s
HEADING_RE = re.compile(r"^(?P<hashes>#{1,6})\s+(?P<text>.*)$")
# Unordered list item
UL_RE = re.compile(r"^\s*[-*+]\s+(?P<text>.*)$")
# Ordered list item
OL_RE = re.compile(r"^\s*\d+[.)]\s+(?P<text>.*)$")
# Table row (starts and is mostly pipes)
TABLE_ROW_RE = re.compile(r"^\s*\|(.+)\|\s*$")
# Table separator row: | --- | :--: | etc.
TABLE_SEP_RE = re.compile(r"^\s*\|?[\s:\-|]+\|?\s*$")
# Inline bold/italic markers we strip to plain text (python-docx runs would be
# overkill for the fast path; keep it readable and never crash).
BOLD_RE = re.compile(r"\*\*(.+?)\*\*")
ITALIC_RE = re.compile(r"(?<!\*)\*(?!\*)(.+?)(?<!\*)\*(?!\*)")
CODE_RE = re.compile(r"`([^`]+)`")


def strip_inline(text):
    """Reduce inline markdown emphasis/code to plain text (best-effort)."""
    text = BOLD_RE.sub(r"\1", text)
    text = ITALIC_RE.sub(r"\1", text)
    text = CODE_RE.sub(r"\1", text)
    return text.strip()


def split_table_cells(line):
    """Split a markdown table row into trimmed cells."""
    inner = line.strip()
    if inner.startswith("|"):
        inner = inner[1:]
    if inner.endswith("|"):
        inner = inner[:-1]
    return [strip_inline(c.strip()) for c in inner.split("|")]


def is_table_separator(line):
    if "|" not in line:
        return False
    body = line.replace("|", "").replace(":", "").replace("-", "").strip()
    return body == "" and "-" in line


def add_table(doc, rows):
    """rows: list of cell-lists. First row treated as header if a separator was seen."""
    if not rows:
        return
    ncols = max(len(r) for r in rows)
    table = doc.add_table(rows=0, cols=ncols)
    try:
        table.style = "Table Grid"
    except Exception:
        pass
    for r in rows:
        cells = table.add_row().cells
        for i in range(ncols):
            cells[i].text = r[i] if i < len(r) else ""


def resolve_image_path(path):
    """Return an absolute path under INPUT_DIR if the referenced image exists."""
    candidates = []
    if os.path.isabs(path):
        candidates.append(path)
    candidates.append(os.path.join(INPUT_DIR, os.path.basename(path)))
    candidates.append(os.path.join(INPUT_DIR, path))
    for c in candidates:
        if os.path.isfile(c):
            return c
    return None


def add_image(doc, alt, path):
    resolved = resolve_image_path(path)
    if resolved is None:
        # Image not available — fall back to a caption so content is not lost.
        if alt:
            doc.add_paragraph(alt)
        return
    try:
        doc.add_picture(resolved, width=Inches(5.5))
    except Exception:
        if alt:
            doc.add_paragraph(alt)


def render(doc, md_text):
    lines = md_text.replace("\r\n", "\n").replace("\r", "\n").split("\n")
    table_buf = []  # accumulated table cell-rows

    def flush_table():
        if table_buf:
            add_table(doc, list(table_buf))
            table_buf.clear()

    for raw in lines:
        line = raw.rstrip()

        # Table handling: accumulate consecutive table rows, skip separators.
        if TABLE_ROW_RE.match(line):
            if is_table_separator(line):
                continue
            table_buf.append(split_table_cells(line))
            continue
        else:
            flush_table()

        if line.strip() == "":
            continue

        m = IMG_RE.match(line.strip())
        if m:
            add_image(doc, m.group("alt"), m.group("path"))
            continue

        m = HEADING_RE.match(line)
        if m:
            level = min(len(m.group("hashes")), 6)
            doc.add_heading(strip_inline(m.group("text")), level=level)
            continue

        m = OL_RE.match(line)
        if m:
            try:
                doc.add_paragraph(strip_inline(m.group("text")), style="List Number")
            except Exception:
                doc.add_paragraph(strip_inline(m.group("text")))
            continue

        m = UL_RE.match(line)
        if m:
            try:
                doc.add_paragraph(strip_inline(m.group("text")), style="List Bullet")
            except Exception:
                doc.add_paragraph(strip_inline(m.group("text")))
            continue

        # Plain paragraph.
        doc.add_paragraph(strip_inline(line))

    flush_table()


def main():
    out_name = sys.argv[1] if len(sys.argv) > 1 else "document.docx"
    if not out_name.lower().endswith(".docx"):
        out_name += ".docx"

    try:
        with open(INPUT_MD, "r", encoding="utf-8") as f:
            md_text = f.read()
    except Exception as e:
        md_text = ""
        sys.stderr.write("md_to_docx: failed to read source.md: %s\n" % e)

    doc = Document()
    try:
        render(doc, md_text)
    except Exception as e:
        # Never crash the sandbox run on a parse error: emit what we have plus a note.
        sys.stderr.write("md_to_docx: render error (rendered best-effort): %s\n" % e)
        doc.add_paragraph(md_text)

    os.makedirs(OUTPUT_DIR, exist_ok=True)
    out_path = os.path.join(OUTPUT_DIR, out_name)
    doc.save(out_path)
    sys.stdout.write("wrote %s\n" % out_path)


if __name__ == "__main__":
    main()

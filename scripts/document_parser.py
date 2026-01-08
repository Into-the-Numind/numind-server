import fitz  # PyMuPDF
import sys
import json
import os
import subprocess
import re
from docx import Document

def clean_text(text):
    """通用的文本清洗函数，去除冗余的空格、空行，仅保留必要的段落结构"""
    if not text:
        return ""
    
    # 1. 规范化换行符
    text = text.replace('\r\n', '\n').replace('\r', '\n')
    
    # 2. 按行处理
    lines = text.split('\n')
    cleaned_lines = []
    
    for line in lines:
        # 清理行内多余空格：将多个连续空格/制表符转换为单个空格
        line = re.sub(r'[ \t]+', ' ', line)
        line = line.strip()
        
        # 过滤掉完全无意义的行（如只有个别特殊符号的行）
        if line:
            cleaned_lines.append(line)
            
    # 3. 重新组合
    # 使用双换行保留基本的段落感，但通过正则确保不会出现 3 个以上的连续换行
    text = '\n\n'.join(cleaned_lines)
    text = re.sub(r'\n{3,}', '\n\n', text)
    
    return text.strip()

def extract_pdf_content(pdf_path):
    try:
        doc = fitz.open(pdf_path)
        full_text = []
        for page_index in range(len(doc)):
            page = doc[page_index]
            blocks = page.get_text("blocks")
            page_content = []
            for b in blocks:
                if b[6] == 0:  # Text block
                    text = b[4].strip()
                    if text:
                        # 块内清理换行，保持语义连贯
                        clean_line = text.replace("\n", " ")
                        page_content.append(clean_line)
            if page_content:
                full_text.append("\n\n".join(page_content))
        doc.close()
        
        raw_content = "\n\n---\n\n".join(full_text)
        return {"success": True, "content": clean_text(raw_content)}
    except Exception as e:
        return {"success": False, "error": f"PDF extraction failed: {str(e)}"}

def extract_docx_content(docx_path):
    try:
        doc = Document(docx_path)
        full_text = []
        # 1. 提取段落
        for para in doc.paragraphs:
            if para.text.strip():
                full_text.append(para.text.strip())
        
        # 2. 提取表格
        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text.strip() for cell in row.cells]
                row_content = " | ".join([c for c in row_text if c])
                if row_content:
                    full_text.append(row_content)
                
        return {"success": True, "content": clean_text("\n\n".join(full_text))}
    except Exception as e:
        return {"success": False, "error": f"DOCX extraction failed: {str(e)}"}

def extract_doc_content(doc_path):
    try:
        result = subprocess.run(['antiword', '-w', '0', doc_path], 
                               capture_output=True, text=True, encoding='utf-8', errors='ignore')
        
        if result.returncode != 0:
            return {"success": False, "error": f"Antiword failed: {result.stderr}"}
        
        content = result.stdout.strip()
        if not content:
            return {"success": False, "error": "No text extracted from .doc file"}
            
        return {"success": True, "content": clean_text(content)}
    except FileNotFoundError:
        return {"success": False, "error": "antiword command not found"}
    except Exception as e:
        return {"success": False, "error": f"DOC extraction failed: {str(e)}"}

def extract_txt_content(txt_path):
    try:
        with open(txt_path, 'r', encoding='utf-8', errors='ignore') as f:
            content = f.read()
        return {"success": True, "content": clean_text(content)}
    except Exception as e:
        return {"success": False, "error": f"TXT extraction failed: {str(e)}"}

if __name__ == "__main__":
    if len(sys.argv) < 2:
        print(json.dumps({"success": False, "error": "No file path provided"}))
        sys.exit(1)
        
    file_path = sys.argv[1]
    ext = os.path.splitext(file_path)[1].lower()
    
    if ext == ".pdf":
        result = extract_pdf_content(file_path)
    elif ext == ".docx":
        result = extract_docx_content(file_path)
    elif ext == ".doc":
        result = extract_doc_content(file_path)
    elif ext in [".txt", ".md"]:
        result = extract_txt_content(file_path)
    else:
        result = {"success": False, "error": f"Unsupported extension: {ext}"}
        
    print(json.dumps(result))
